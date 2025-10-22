"""
Document Processing Module
Handles document loading, text extraction, and chunking for various file formats
"""

import os
from typing import List, Dict, Any
from io import BytesIO

import PyPDF2
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangChainDocument
import re

from config import CHUNK_SIZE, CHUNK_OVERLAP, SUPPORTED_FORMATS, MAX_FILE_SIZE_BYTES


class DocumentProcessor:
    """Handles document processing operations"""
    
    def __init__(self):
        """Initialize the document processor with text splitter"""
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def validate_file(self, file_name: str, file_size: int) -> tuple[bool, str]:
        """
        Validate file format and size
        
        Args:
            file_name: Name of the file
            file_size: Size of the file in bytes
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        # Check file extension
        file_ext = os.path.splitext(file_name)[1].lower()
        if file_ext not in SUPPORTED_FORMATS:
            return False, f"Unsupported file format. Supported formats: {', '.join(SUPPORTED_FORMATS)}"
        
        # Check file size
        if file_size > MAX_FILE_SIZE_BYTES:
            return False, f"File size exceeds {MAX_FILE_SIZE_BYTES / (1024*1024)}MB limit"
        
        return True, ""
    
    def _clean_text(self, text: str) -> str:
        """
        Normalize whitespace and fix broken line wraps from PDF extraction.
        - Collapse single newlines into spaces while preserving paragraph breaks
        - Remove excessive spaces
        - Fix hyphenated line breaks
        """
        if not text:
            return ""
        # Normalize line endings
        text = text.replace("\r", "\n")
        # Fix hyphenated line breaks: "exam-\nple" -> "example"
        text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
        # Remove page markers that add noise
        text = re.sub(r"\s*--- Page \d+ ---\s*", "\n", text)
        # Replace single newlines (not part of a blank-line paragraph break) with spaces
        text = re.sub(r"(?<!\n)\n(?!\n)", " ", text)
        # Collapse 3+ newlines to 2 (paragraph break)
        text = re.sub(r"\n{3,}", "\n\n", text)
        # Collapse repeated spaces
        text = re.sub(r"[ \t]{2,}", " ", text)
        return text.strip()

    def extract_text_from_pdf(self, file_bytes: bytes) -> tuple[str, Dict[str, Any]]:
        """
        Extract text from PDF file
        
        Args:
            file_bytes: PDF file content as bytes
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        text = ""
        metadata = {"pages": 0}
        
        try:
            pdf_file = BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            metadata["pages"] = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}"
            
        except Exception as e:
            raise ValueError(f"Error extracting text from PDF: {str(e)}")
        
        # Clean up broken formatting common in PDFs
        text = self._clean_text(text)
        return text, metadata
    
    def extract_text_from_docx(self, file_bytes: bytes) -> tuple[str, Dict[str, Any]]:
        """
        Extract text from DOCX file
        
        Args:
            file_bytes: DOCX file content as bytes
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        text = ""
        metadata = {"paragraphs": 0}
        
        try:
            docx_file = BytesIO(file_bytes)
            doc = Document(docx_file)
            
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            metadata["paragraphs"] = len(paragraphs)
            text = "\n\n".join(paragraphs)
            
        except Exception as e:
            raise ValueError(f"Error extracting text from DOCX: {str(e)}")
        
        text = self._clean_text(text)
        return text, metadata
    
    def extract_text_from_txt(self, file_bytes: bytes) -> tuple[str, Dict[str, Any]]:
        """
        Extract text from TXT file
        
        Args:
            file_bytes: TXT file content as bytes
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        text = ""
        metadata = {}
        
        try:
            text = file_bytes.decode('utf-8')
            metadata["lines"] = len(text.split('\n'))
        except UnicodeDecodeError:
            try:
                text = file_bytes.decode('latin-1')
                metadata["lines"] = len(text.split('\n'))
            except Exception as e:
                raise ValueError(f"Error decoding text file: {str(e)}")
        
        text = self._clean_text(text)
        return text, metadata
    
    def process_document(self, file_name: str, file_bytes: bytes) -> List[LangChainDocument]:
        """
        Process a document: extract text, chunk it, and create LangChain documents
        
        Args:
            file_name: Name of the file
            file_bytes: File content as bytes
            
        Returns:
            List of LangChain Document objects with chunks and metadata
        """
        # Validate file
        is_valid, error_msg = self.validate_file(file_name, len(file_bytes))
        if not is_valid:
            raise ValueError(error_msg)
        
        # Extract text based on file type
        file_ext = os.path.splitext(file_name)[1].lower()
        
        if file_ext == ".pdf":
            text, doc_metadata = self.extract_text_from_pdf(file_bytes)
        elif file_ext == ".docx":
            text, doc_metadata = self.extract_text_from_docx(file_bytes)
        elif file_ext == ".txt":
            text, doc_metadata = self.extract_text_from_txt(file_bytes)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        # Check if text was extracted
        if not text or not text.strip():
            raise ValueError("No text could be extracted from the document")
        
        # Split text into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Create LangChain documents with metadata
        documents = []
        for idx, chunk in enumerate(chunks):
            metadata = {
                "source": file_name,
                "chunk_index": idx,
                "total_chunks": len(chunks),
                **doc_metadata
            }
            documents.append(LangChainDocument(page_content=chunk, metadata=metadata))
        
        return documents
    
    def get_document_info(self, documents: List[LangChainDocument]) -> Dict[str, Any]:
        """
        Get summary information about processed documents
        
        Args:
            documents: List of LangChain documents
            
        Returns:
            Dictionary with document statistics
        """
        if not documents:
            return {"total_chunks": 0, "total_characters": 0}
        
        total_chars = sum(len(doc.page_content) for doc in documents)
        
        return {
            "source": documents[0].metadata.get("source", "Unknown"),
            "total_chunks": len(documents),
            "total_characters": total_chars,
            "metadata": {k: v for k, v in documents[0].metadata.items() 
                        if k not in ["chunk_index", "total_chunks"]}
        }
